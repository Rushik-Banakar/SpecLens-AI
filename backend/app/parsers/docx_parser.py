"""DOCX text extraction using python-docx."""

from __future__ import annotations

import docx


def parse_docx(file_path: str) -> str:
    """Extract paragraph and table text from a DOCX document.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Newline-joined document text.

    Raises:
        ValueError: If the DOCX cannot be opened or parsed.
    """
    try:
        document = docx.Document(file_path)
        full_text: list[str] = []
        for paragraph in document.paragraphs:
            full_text.append(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        return "\n".join(full_text)
    except Exception as exc:
        raise ValueError(f"Failed to extract text from DOCX: {exc}") from exc
